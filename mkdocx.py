# -*- coding: utf-8 -*-
"""대화문 묶음 → 검토용 docx. 기존 다섯 편과 같은 꼴로 찍는다."""
import json, io, sys, hashlib
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1]; OUT = sys.argv[2]; TITLE = sys.argv[3]
d = json.load(io.open(SRC, encoding="utf-8"))

FONT = "맑은 고딕"
def setfont(run, size, bold=False, name=FONT, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color: run.font.color.rgb = RGBColor.from_string(color)

def para(doc, text, size, bold=False, name=FONT, space=2, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    setfont(p.add_run(text), size, bold, name)
    return p

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd"); el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto"); el.set(qn("w:fill"), fill)
    tcPr.append(el)

def borders(tbl):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for e in ("top","left","bottom","right","insideH","insideV"):
        x = OxmlElement("w:"+e); x.set(qn("w:val"),"single")
        x.set(qn("w:color"),"auto"); x.set(qn("w:sz"),"4"); b.append(x)
    tblPr.append(b)

def cell_text(cell, text, size, bold=False, name=FONT):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    setfont(p.add_run(text), size, bold, name)

doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Cm(21.0), Cm(29.7)   # A4
s.left_margin = s.right_margin = Cm(1.27)
s.top_margin = s.bottom_margin = Cm(1.2)

def widths(tbl, *cm):
    """표 열 폭.
       ★ cell.width 만 주면 안 먹는다. 워드는 **tblGrid** 를 먼저 보고,
         거기 값이 없으면 남은 폭을 열 개수로 나눠 버린다(그래서 반반이 됐다).
         tblGrid·tblW·각 셀 tcW 를 다 맞춰야 한다."""
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")): tblPr.remove(old)
    w = OxmlElement("w:tblW"); w.set(qn("w:type"), "dxa")
    w.set(qn("w:w"), str(int(sum(cm) * 567))); tblPr.append(w)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None: tbl._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for x in cm:
        g = OxmlElement("w:gridCol"); g.set(qn("w:w"), str(int(x * 567))); grid.append(g)
    tbl._tbl.insert(list(tbl._tbl).index(tblPr) + 1, grid)
    for row in tbl.rows:
        for c, x in zip(row.cells, cm):
            c.width = Cm(x)

h = doc.add_heading(f"〈검토용〉 상호작용 대화 능력 학습 대화문 — {TITLE}", level=1)
setfont(h.runs[0], 11, True)
para(doc, d["gist"], 7.5)
para(doc, "자리는 〈표12 최종 주제 등급화〉의 2·3급 주제에서만 골랐습니다. 중급 학습자에게 친숙한 주제란 이미 배운 주제입니다 — 요소를 가르치려다 낯선 주제로 학습자를 막지 않기 위해서입니다.", 7.0)
para(doc, "각 편은 자리에 맞는 화계 하나로 짓고, 다섯 편 안에서 세 화계를 흩어 놓았습니다. 학습자는 ⓪들어가기 → ①듣기 → ②추측 → ③뜻풀이 → ④문형 → ⑤연습 순으로 만납니다.", 7.5, space=3)
if d["items"] and d["items"][0].get("quiz_type") == "order":
    para(doc, "※ 이 묶음만 ②추측이 다릅니다 — 세 갈래 고르기가 아니라 조각을 차례대로 눌러 쌓는 방식입니다. 한 발화가 아니라 대화 한 판의 흐름을 보는 것이라, 세 갈래로는 배울 것이 남지 않습니다.", 7.0, space=2)
    para(doc, "※ 다섯 편 모두 복합 대화입니다. 단순 대화(길 묻기 등)는 하나의 대화이동 연속체로 완결되므로 기능 단계가 없습니다(박용익 2014).", 7.0, space=2)
    para(doc, "※ 단계 이름을 두 벌 두었습니다. 앞의 것은 학습자가 보는 말로 연구 용어를 쓰지 않았고, 〔  〕 안의 것은 논문에 쓰는 말로 화면에는 나오지 않습니다. 근거는 박용익(2014) 〈복합 대화와 그 기능 단계〉와 이남호·이찬규(2025)의 구매 대화 원형입니다.", 7.0, space=8)

for n, it in enumerate(d["items"], 1):
    h2 = doc.add_heading(f"{n}. {it['sub']}   {it['place']}  ·  {it['rel']}", level=2)
    setfont(h2.runs[0], 9.5, True)
    _lv = f"〈표12〉 {it['topic_lv']}"
    if it.get("acad_type"):
        _lv += f"      ·  대화 유형: {it['acad_type']}"
    para(doc, _lv, 6.5, space=3)

    # ── ① 대화문 ──
    t = doc.add_table(rows=0, cols=2); borders(t)
    r = t.add_row()
    cell_text(r.cells[0], "", 7.5); shade(r.cells[0], "E8E3DA")
    cell_text(r.cells[1], "대화문  ( ▶ 표시한 줄이 배울 자리 )", 7.5, True); shade(r.cells[1], "E8E3DA")
    for i, l in enumerate(it["script"]):
        r = t.add_row()
        cell_text(r.cells[0], "나" if l["speaker"] == "user" else "호아랑", 6.5)
        mk = i == it["mark"]
        cell_text(r.cells[1], ("▶  " if mk else "") + l["text"], 7.5, mk)
        if mk: shade(r.cells[0], "FFF6DE"); shade(r.cells[1], "FFF6DE")
    widths(t, 1.59, 15.42)
    para(doc, "", 4, space=2)

    # ── ② 추측 ──
    q = it["quiz"]
    t2 = doc.add_table(rows=0, cols=1); borders(t2)
    r = t2.add_row()
    cell_text(r.cells[0], "②  " + q["q"], 7.5, True); shade(r.cells[0], "EEF2F8")

    if it.get("quiz_type") == "order":
        # ★ 차례 맞히기 (v136) — 세 갈래가 아니라 조각을 순서대로 쌓는다.
        #   검토할 때는 **맞는 차례**와 그 단계가 어느 줄에서 시작하는지가 보여야 한다.
        # ★ 두 이름을 나란히 찍는다.
        #   왼쪽 = 학습자가 보는 말 (연구 용어 없음)
        #   오른쪽 = 논문에 쓰는 말 (화면에는 안 나온다)
        for n, stg in enumerate(it.get("stages") or [], 1):
            r = t2.add_row()
            c = r.cells[0]
            pp = c.paragraphs[0]
            pp.paragraph_format.space_before = Pt(1); pp.paragraph_format.space_after = Pt(1)
            setfont(pp.add_run(f"{n})  {stg['name']}"), 7.0, True)
            setfont(pp.add_run(f"        〔{stg.get('acad', '')}〕"), 6.5, False, color="7A6A55")
            setfont(pp.add_run(f"    ← {stg['at'] + 1}번째 줄부터"), 6.0, False, color="9A8C78")
    else:
        # ★ 자리를 섞는다. 검토용이라 「← 정답」은 남기되 늘 ⓐ에 놓지는 않는다.
        #   정답이 언제나 첫 줄이면 검토하는 사람의 눈이 거기부터 가서
        #   「오답 둘이 정말 그럴듯한가」를 제대로 못 본다.
        #   앱은 이미 os.urandom 으로 섞는다(main.py) — 학습자 화면과 어긋나지 않는다.
        #   여기서는 **id 를 해시해 자리를 정한다** — 다시 찍어도 같은 자리라야
        #   종이에 적어 둔 검토 의견과 어긋나지 않는다.
        #   (글자 합만 쓰면 repair-1·listen-1·move-1 이 죄다 같은 자리가 된다)
        keys = ["right", "wrong1", "wrong2"]
        seed = hashlib.md5(it["id"].encode("utf-8")).digest()[0]
        keys = keys[seed % 3:] + keys[:seed % 3]
        for mark, key in zip(("ⓐ", "ⓑ", "ⓒ"), keys):
            r = t2.add_row()
            right = key == "right"
            cell_text(r.cells[0], f"{mark} {q[key]}" + ("        ← 정답" if right else ""), 7.0, right)

    r = t2.add_row()
    cell_text(r.cells[0], "틀렸을 때 →  " + q.get("hint", ""), 6.5); shade(r.cells[0], "FAF8F4")
    widths(t2, 17.0)
    para(doc, "", 4, space=2)

    # ── ③ 뜻풀이 ──
    para(doc, "③  뜻풀이", 7.5, True, space=1)
    for i, m in enumerate(it["meaning"], 1):
        para(doc, f"{i})  {m}", 7.0, space=1)
    # ── ④ 문형 ──
    # ★ 화면에는 「~」인 채로가 아니라 **그 대화문의 말로 채워져** 뜬다 (v138).
    #   검토도 화면에서 보이는 대로 해야 하므로 채운 꼴을 함께 찍는다.
    #   〔  〕 안이 갈아 끼우는 자리(화면에서는 검은색), 밖이 외울 뼈대(붉은색).
    para(doc, "④  문형", 7.5, True, space=1)
    _ff = it.get("forms_filled") or []
    for i, x in enumerate(it["forms"]):
        parts = _ff[i] if i < len(_ff) else [[x, 1]]
        shown = "".join(("〔" + t + "〕") if not c else t for t, c in parts)
        para(doc, f"   ·  {x}" + (f"      →   {shown}" if "~" in x else ""),
             7.0, name="GmarketSansLight", space=0)
    para(doc, "", 4, space=2)
    # ── ⑤ 연습 ──
    para(doc, "⑤  연습", 7.5, True, space=1)
    for i, dr in enumerate(it["drills"], 1):
        para(doc, f"   {i})  호아랑: {dr['cue']}", 6.5, space=0)
        para(doc, f"        나: {dr['text']}", 7.0, True, space=0)
        if dr.get("follow"):
            para(doc, f"        호아랑: {dr['follow']}", 6.5, space=1)
    para(doc, "", 6, space=6)

doc.save(OUT)
print(f"✅ {OUT}")
